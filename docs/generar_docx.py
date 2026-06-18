# -*- coding: utf-8 -*-
"""
Genera el documento Word con los valores VERIFICADOS de calidad del gas natural
para el proyecto ENAGÁS — Chatbot de comparación regulatoria de calidad de gas
natural en Europa (revisión 2026-06).

DATA-DRIVEN: lee directamente la ontología verificada como única fuente de verdad.
    Fuente:  ../data/ontologia/ontologia_enagas.yaml
    Salida:  ./Calidad_Gas_Valores_Verificados.docx

Uso:
    python -m pip install python-docx pyyaml
    python ENAG-S_G1/docs/generar_docx.py
"""

import os
import yaml

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------------- #
# Rutas
# --------------------------------------------------------------------------- #
HERE = os.path.dirname(os.path.abspath(__file__))
ONTOLOGIA = os.path.join(HERE, "..", "data", "ontologia", "ontologia_enagas.yaml")
SALIDA = os.path.join(HERE, "Calidad_Gas_Valores_Verificados.docx")

JURISDICCIONES = ["ES", "PT", "FR", "UE"]
JUR_NOMBRE = {"ES": "España (ES)", "PT": "Portugal (PT)",
              "FR": "Francia (FR)", "UE": "UE"}

UNIT_DISPLAY = {
    "kWh_per_nm3": "kWh/m³", "MJ_per_nm3": "MJ/m³", "mg_per_nm3": "mg/m³",
    "mg_per_sm3": "mg/m³(15ºC)", "pct_mol": "% mol", "pct_vol": "% vol",
    "ppm_vol": "ppm", "grados_C": "ºC", "adimensional": "", None: "",
}

# --------------------------------------------------------------------------- #
# Paleta y utilidades de estilo
# --------------------------------------------------------------------------- #
AZUL_ENAGAS = RGBColor(0x00, 0x3D, 0x6B)
AZUL_CLARO = RGBColor(0x1F, 0x6F, 0xB2)
GRIS = RGBColor(0x55, 0x55, 0x55)
GRIS_CLARO = RGBColor(0x99, 0x99, 0x99)
VERDE = RGBColor(0x1E, 0x7A, 0x34)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9, align="left", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def style_header_row(row, fill="003D6B"):
    for cell in row.cells:
        set_cell_background(cell, fill)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = BLANCO
                r.bold = True


# --------------------------------------------------------------------------- #
# Formateo de valores desde la ontología
# --------------------------------------------------------------------------- #
def fmt_num(x):
    return ("%g" % x).replace(".", ",")


def format_limite(lim):
    """Devuelve (texto, estado) para mostrar en una celda."""
    estado = lim.get("estado_verificacion")
    if estado == "NO_VERIFICABLE_SIN_FUENTE":
        return "No fija", estado
    if estado == "PENDIENTE_EXTRACCION":
        return "Pendiente", estado
    unidad = UNIT_DISPLAY.get(lim.get("unidad"), lim.get("unidad") or "")
    if lim.get("tipo_limite") == "rango":
        vmin, vmax = lim.get("valor_min"), lim.get("valor_max")
        if vmin is None or vmax is None:
            return "No fija", "NO_VERIFICABLE_SIN_FUENTE"
        base = f"{fmt_num(vmin)} – {fmt_num(vmax)}"
    else:
        v = lim.get("valor")
        if v is None:
            return "No fija", "NO_VERIFICABLE_SIN_FUENTE"
        base = f"≤ {fmt_num(v)}"
    txt = (base + " " + unidad).strip()
    p = lim.get("presion_referencia_bar")
    if p is not None:
        txt += f" @ {p} bar"
    return txt, estado


def color_estado(estado):
    return VERDE if estado == "VERIFICADO" else GRIS_CLARO


# --------------------------------------------------------------------------- #
# Construcción del documento
# --------------------------------------------------------------------------- #
def build(onto):
    o = onto["ontologia"]
    params = onto["parametros"]
    fuentes = o["fuentes_normativas"]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ---- Portada ----
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(t, "Especificaciones de Calidad del Gas Natural", bold=True, color=AZUL_ENAGAS, size=20)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(t2, "Valores Verificados — Comparación multinacional", bold=True, color=AZUL_CLARO, size=15)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(s, f"Conjunto de datos verificado — revisión {o.get('fecha_revision','')}  ·  ontología v{o.get('version','')}",
            italic=True, color=GRIS, size=11)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(s2, "Proyecto ENAGÁS — Calidad de gas natural en Europa (ES · PT · FR · UE)",
            italic=True, color=GRIS, size=11)
    doc.add_paragraph()

    # ---- 1. Alcance ----
    doc.add_heading("1. Alcance y metodología", level=1)
    p = doc.add_paragraph()
    add_run(p, "El proyecto compara los 10 parámetros de calidad del gas natural del "
               "alcance (excluye Polvo/Partículas) entre España, Portugal, Francia y el "
               "marco europeo. Cada valor de esta tabla está contrastado verbatim contra "
               "su documento oficial (artículo/tabla + página). Cuando una norma no fija "
               "un parámetro, se indica «No fija» y no se rellena con ninguna cifra.")
    p = doc.add_paragraph()
    add_run(p, "Cobertura verificada: ", bold=True)
    add_run(p, "ES 10/10 · FR 10/10 · PT 6/10 · UE 0/10 (el Reglamento (UE) 2015/703 no "
               "fija límites numéricos de calidad).")

    # ---- 2. Condiciones de referencia ----
    doc.add_heading("2. Condiciones de referencia por jurisdicción", level=1)
    jur = {j["codigo"]: j for j in o["jurisdicciones"]}
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    set_cell_text(tbl.rows[0].cells[0], "Jurisdicción", bold=True)
    set_cell_text(tbl.rows[0].cells[1], "Condiciones de referencia", bold=True)
    style_header_row(tbl.rows[0])
    for cod in JURISDICCIONES:
        row = tbl.add_row().cells
        set_cell_text(row[0], jur[cod]["nombre"], bold=True, color=AZUL_ENAGAS)
        set_cell_text(row[1], str(jur[cod].get("condiciones_referencia_defecto", "")))

    # ---- 3. Tabla principal ----
    doc.add_heading("3. Valores verificados (parámetro × jurisdicción)", level=1)
    tbl = doc.add_table(rows=1, cols=1 + len(JURISDICCIONES))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    set_cell_text(hdr[0], "Parámetro", bold=True)
    for i, cod in enumerate(JURISDICCIONES, start=1):
        set_cell_text(hdr[i], JUR_NOMBRE[cod], bold=True, align="center")
    style_header_row(tbl.rows[0])

    for pid, p in params.items():
        nombre = p.get("nombre_completo", pid)
        simbolo = p.get("simbolo", "")
        # Añade el símbolo solo si es una abreviatura corta (IW, PCS, O₂, CO₂…)
        etiqueta = nombre + (f" ({simbolo})" if simbolo and len(simbolo) <= 8 and simbolo not in nombre else "")
        cells = tbl.add_row().cells
        set_cell_text(cells[0], etiqueta, bold=True, color=AZUL_ENAGAS, size=9)
        for i, cod in enumerate(JURISDICCIONES, start=1):
            txt, estado = format_limite(p["limites"][cod])
            set_cell_text(cells[i], txt, align="center", size=9, color=color_estado(estado))

    p = doc.add_paragraph()
    add_run(p, "Verde = VERIFICADO verbatim.  Gris = «No fija» (NO_VERIFICABLE_SIN_FUENTE).",
            italic=True, color=GRIS, size=8.5)

    # ---- 4. Fuentes ----
    doc.add_heading("4. Fuentes oficiales", level=1)
    for f in fuentes:
        if f.get("tipo") in ("buenas_practicas", "norma_tecnica"):
            continue
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, f"[{f.get('pais','')}] {f.get('nombre','')}", bold=True, size=9.5)
        extra = f.get("publicacion") or f.get("tabla_calidad")
        if extra:
            add_run(p, f" — {extra}", size=9, color=GRIS)

    # ---- Cierre ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p, f"Documento generado automáticamente desde la ontología verificada "
               f"(ontologia_enagas.yaml v{o.get('version','')}). Ninguna cifra inventada.",
            italic=True, color=GRIS, size=9)

    return doc


def main():
    with open(ONTOLOGIA, encoding="utf-8") as fh:
        onto = yaml.safe_load(fh)
    doc = build(onto)
    doc.save(SALIDA)
    print(f"OK -> {SALIDA}")


if __name__ == "__main__":
    main()
